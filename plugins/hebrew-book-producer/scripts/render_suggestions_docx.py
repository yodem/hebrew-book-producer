#!/usr/bin/env python3
"""Render a changes.json + chapter markdown into a .docx with tracked changes.

Usage:
    render_suggestions_docx.py \
        --changes <path/to/changes.json> \
        --source <chapters/ dir> \
        --out <out dir>

For each change object whose `file` matches a chapter under <source>:
- Build a new .docx with the chapter's paragraphs.
- For each paragraph that contains the `before` text:
    - Replace `before` with a tracked-change pair (w:del + w:ins).
    - Add a bookmark chg_<change_id> at the change site.
    - Add a comment with the rationale text.

Output: <out>/<chapter>.suggestions.docx
"""
from __future__ import annotations
import argparse
import json
import sys
from pathlib import Path
from typing import Optional, Tuple, List, Dict
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent))
from docx_tracked_changes import (  # noqa: E402
    add_tracked_change,
    add_change_bookmark,
    add_comment,
)


def _read_chapter(chapters_dir: Path, file_field: str) -> Optional[Tuple[Path, List[str]]]:
    name = Path(file_field).name
    path = chapters_dir / name
    if not path.is_file():
        return None
    text = path.read_text(encoding="utf-8")
    paragraphs = [p for p in text.split("\n\n") if p.strip()]
    return path, paragraphs


def _group_changes_by_chapter(changes: List[Dict]) -> Dict[str, List[Dict]]:
    grouped: Dict[str, List[Dict]] = {}
    for c in changes:
        f = c.get("file", "")
        chapter = Path(f).stem  # "ch01"
        grouped.setdefault(chapter, []).append(c)
    return grouped


def render_chapter(
    chapter_id: str,
    paragraphs: List[str],
    changes_for_chapter: List[Dict],
    out_path: Path,
) -> int:
    doc = Document()
    embedded = 0
    comment_counter = [0]

    for paragraph_text in paragraphs:
        applicable: List[Dict] = []
        remaining: List[Dict] = []
        for c in changes_for_chapter:
            before = c.get("before", "")
            if before and before in paragraph_text:
                applicable.append(c)
            else:
                remaining.append(c)
        changes_for_chapter = remaining

        p = doc.add_paragraph()

        if not applicable:
            p.add_run(paragraph_text)
            continue

        c = applicable[0]
        before = c["before"]
        after = c.get("after", "")
        rationale = c.get("rationale", "")
        change_id = c["change_id"]
        author = "hebrew-book-producer"

        idx = paragraph_text.find(before)
        prefix = paragraph_text[:idx]
        suffix = paragraph_text[idx + len(before):]

        if prefix:
            p.add_run(prefix)
        add_tracked_change(p, before=before, after=after, author=author)
        if suffix:
            p.add_run(suffix)
        add_change_bookmark(p, change_id=change_id)
        comment_counter[0] += 1
        try:
            add_comment(
                document=doc,
                paragraph=p,
                author=author,
                text=rationale,
                comment_id=comment_counter[0],
            )
        except Exception as e:
            p.add_run(f"  [הערה: {rationale}]")
            print(f"comment fallback: {e}", file=sys.stderr)

        embedded += 1

        for extra in applicable[1:]:
            print(
                f"skipped multi-change paragraph: {chapter_id} change_id={extra['change_id']}",
                file=sys.stderr,
            )

    for c in changes_for_chapter:
        print(
            f"skipped: no match for change_id={c['change_id']} before='{c.get('before', '')[:30]}...'",
            file=sys.stderr,
        )

    doc.save(out_path)
    return embedded


def main() -> int:
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("--changes", required=True, help="path to changes.json")
    p.add_argument("--source", required=True, help="chapters/ directory")
    p.add_argument("--out", required=True, help="output directory")
    args = p.parse_args()

    changes_path = Path(args.changes)
    source_dir = Path(args.source)
    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    payload = json.loads(changes_path.read_text(encoding="utf-8"))
    changes = payload.get("changes", [])
    grouped = _group_changes_by_chapter(changes)

    total_embedded = 0
    for chapter_id, chapter_changes in grouped.items():
        first_file = chapter_changes[0]["file"]
        resolved = _read_chapter(source_dir, first_file)
        if resolved is None:
            print(f"skipped chapter (file not found): {first_file}", file=sys.stderr)
            continue
        path, paragraphs = resolved
        out_path = out_dir / f"{chapter_id}.suggestions.docx"
        embedded = render_chapter(chapter_id, paragraphs, chapter_changes, out_path)
        total_embedded += embedded
        print(f"{chapter_id}: embedded {embedded} changes -> {out_path}")

    print(f"total embedded: {total_embedded}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
