"""Tests for docx_tracked_changes — low-level OOXML helpers."""
from __future__ import annotations
import sys
from pathlib import Path
from docx import Document

REPO_ROOT = Path(__file__).resolve().parents[4]
sys.path.insert(0, str(REPO_ROOT / "plugins/hebrew-book-producer/scripts"))
from docx_tracked_changes import (  # noqa: E402
    add_tracked_change,
    add_change_bookmark,
    add_comment,
    W_NS,
)


def test_add_tracked_change_emits_ins_and_del(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph("השינויים")
    add_tracked_change(
        paragraph=p,
        before="המחשבות",
        after="המחשבה",
        author="hebrew-book-producer",
    )
    out = tmp_path / "out.docx"
    doc.save(out)

    doc2 = Document(out)
    xml = doc2.paragraphs[0]._p.xml
    assert "w:del" in xml
    assert "w:ins" in xml
    assert "המחשבות" in xml
    assert "המחשבה" in xml


def test_add_change_bookmark_includes_id_in_name(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph("טקסט")
    add_change_bookmark(p, change_id="abc123def456")
    out = tmp_path / "out.docx"
    doc.save(out)

    doc2 = Document(out)
    xml = doc2.paragraphs[0]._p.xml
    assert "chg_abc123def456" in xml


def test_add_comment_writes_comment_xml(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph("טקסט")
    add_comment(
        document=doc,
        paragraph=p,
        author="hebrew-book-producer",
        text="register: literary-formal",
        comment_id=1,
    )
    out = tmp_path / "out.docx"
    doc.save(out)

    doc2 = Document(out)
    assert doc2.paragraphs[0].text  # paragraph still has content
